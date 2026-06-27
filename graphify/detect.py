# file discovery, type classification, and corpus health checks
from __future__ import annotations
import fnmatch
import json
import os
import re
import shlex
from concurrent.futures import ThreadPoolExecutor
from enum import Enum
from pathlib import Path

from graphify.google_workspace import (
    GOOGLE_WORKSPACE_EXTENSIONS,
    convert_google_workspace_file,
    google_workspace_enabled,
)
from graphify.paths import GRAPHIFY_OUT, GRAPHIFY_OUT_NAME, out_path


class FileType(str, Enum):
    CODE = "code"
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"
    VIDEO = "video"


_MANIFEST_PATH = str(out_path("manifest.json"))

CODE_EXTENSIONS = {'.py', '.ts', '.tsx', '.js', '.jsx', '.mjs', '.ejs', '.ets', '.go', '.rs', '.java', '.groovy', '.gradle', '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp', '.cu', '.cuh', '.metal', '.rb', '.swift', '.kt', '.kts', '.cs', '.scala', '.php', '.lua', '.luau', '.toc', '.zig', '.ps1', '.psm1', '.psd1', '.ex', '.exs', '.m', '.mm', '.jl', '.vue', '.svelte', '.astro', '.dart', '.v', '.sv', '.svh', '.sql', '.r', '.f', '.F', '.f90', '.F90', '.f95', '.F95', '.f03', '.F03', '.f08', '.F08', '.pas', '.pp', '.dpr', '.dpk', '.lpr', '.inc', '.dfm', '.lfm', '.lpk', '.sh', '.bash', '.json', '.tf', '.tfvars', '.hcl', '.dm', '.dme', '.dmi', '.dmm', '.dmf', '.sln', '.slnx', '.csproj', '.fsproj', '.vbproj', '.xaml', '.razor', '.cshtml', '.cls', '.trigger'}
DOC_EXTENSIONS = {'.md', '.mdx', '.qmd', '.txt', '.rst', '.html', '.yaml', '.yml'}
PAPER_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'}
OFFICE_EXTENSIONS = {'.docx', '.xlsx'}
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.webm', '.mkv', '.avi', '.m4v', '.mp3', '.wav', '.m4a', '.ogg'}

CORPUS_WARN_THRESHOLD = 50_000    # words - below this, warn "you may not need a graph"
CORPUS_UPPER_THRESHOLD = 500_000  # words - above this, warn about token cost
FILE_COUNT_UPPER = 500             # files - above this, warn about token cost

# Resource caps for parsing untrusted office/PDF files (F2). A corpus is
# attacker-controllable (graphify runs on cloned/shared folders), and .docx/.xlsx
# are zip+XML containers: a few-KB zip-bomb can decompress to gigabytes and
# OOM-kill the process at load_workbook/Document time. Screen the file before any
# parser touches it.
_OFFICE_MAX_RAW_BYTES = 50 * 1024 * 1024            # 50 MiB on-disk
_OFFICE_MAX_DECOMPRESSED_BYTES = 512 * 1024 * 1024  # 512 MiB total uncompressed
_OFFICE_MAX_COMPRESSION_RATIO = 200                 # uncompressed : compressed


def _file_within_size_cap(path: Path, cap: int = _OFFICE_MAX_RAW_BYTES) -> bool:
    """True if *path* exists and its on-disk size is within *cap*."""
    try:
        return path.stat().st_size <= cap
    except OSError:
        return False


def _zip_within_caps(path: Path) -> bool:
    """Reject a zip-based office file that is a likely zip/XML bomb.

    Two layers, because the zip central-directory sizes are attacker-controlled:
    1. A cheap pre-filter on the declared sizes (on-disk cap, summed-uncompressed
       cap, compression ratio) that rejects an honest bomb without decompressing.
    2. An authoritative pass that stream-decompresses every member with a hard
       byte ceiling, so a member that under-declares its size in the central
       directory cannot expand past the cap undetected. Decompression is chunked
       and bounded, so checking a bomb never materializes more than the ceiling.
    """
    import zipfile
    if not _file_within_size_cap(path):
        return False
    try:
        with zipfile.ZipFile(path) as zf:
            infos = zf.infolist()
            compressed = sum(i.compress_size for i in infos) or 1
            declared = sum(i.file_size for i in infos)
            if declared > _OFFICE_MAX_DECOMPRESSED_BYTES:
                return False
            if declared / compressed > _OFFICE_MAX_COMPRESSION_RATIO:
                return False
            total = 0
            for info in infos:
                with zf.open(info) as member:
                    while True:
                        chunk = member.read(1024 * 1024)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > _OFFICE_MAX_DECOMPRESSED_BYTES:
                            return False
    except (zipfile.BadZipFile, OSError, EOFError):
        return False
    return True

# Parent directories whose contents are always sensitive.
# Checked against path.parts[:-1] (parents only) so a root-level file named
# "credentials" or "secrets" is not falsely flagged by this stage.
_SENSITIVE_DIRS = frozenset({
    ".ssh", ".gnupg", ".aws", ".gcloud", "secrets", ".secrets", "credentials",
})

# Files that may contain secrets - skip silently. These patterns are specific
# (extensions, exact credential-store names) and always apply.
_SENSITIVE_PATTERNS = [
    re.compile(r'(^|[\\/])\.(env|envrc)(\.|$)', re.IGNORECASE),
    re.compile(r'\.(pem|key|p12|pfx|cert|crt|der|p8)$', re.IGNORECASE),
    re.compile(r'(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$'),
    re.compile(r'(\.netrc|\.pgpass|\.htpasswd)$', re.IGNORECASE),
    re.compile(r'(aws_credentials|gcloud_credentials|service.account)', re.IGNORECASE),
]

# Generic keyword patterns - these only count when the keyword is LOAD-BEARING
# in the filename (see _generic_keyword_hit), because a keyword buried mid-phrase
# in a long descriptive slug names a topic, not a credential store:
# "token-economics-of-recall.md" is a note ABOUT tokens; "api_token.txt" IS one.
# Uses lookarounds instead of \b so underscore-prefixed names like api_token.txt
# match. Both patterns use (?![a-zA-Z]) so that the trailing-underscore behavior
# is consistent: "secret_store.txt" IS flagged, "tokenizer.py" is NOT (because
# "i" after "token" is alpha and blocks the match).
# `token` is kept separate because its longer suffix "izer"/"ize" is the only
# common false-positive; other keywords have no such well-known derivatives.
_GENERIC_KEYWORD_PATTERNS = [
    re.compile(r'(?<![a-zA-Z0-9])(credential|secret|passwd|password|private_key)s?(?![a-zA-Z])', re.IGNORECASE),
    re.compile(r'(?<![a-zA-Z0-9])tokens?(?![a-zA-Z])', re.IGNORECASE),
]

# Word separators for the load-bearing check (underscore intentionally included;
# multi-word keywords like private_key are handled by the end-of-stem check,
# which runs before word counting).
_WORD_SPLIT = re.compile(r'[-_\s]+')


def _generic_keyword_hit(name: str) -> bool:
    """True if a generic secret keyword appears load-bearing in the filename.

    Secret-store files name their contents, and in English compounds the
    content noun is the head, which comes last: "github-personal-access-token",
    "api_token", "oauth_token". A keyword that is neither at the end of the
    stem nor in a short (<=2 word) name is a topic word in a descriptive slug
    ("token-economics-of-recall.md", "password-policy-discussion.md") and must
    not cause the file to be silently dropped from the graph (#436, #718).
    """
    # Stem = name up to the first dot, ignoring leading dots so dotfiles like
    # ".token" keep their keyword ("" stems would never match).
    stem = name.lstrip('.').split('.')[0]
    for pat in _GENERIC_KEYWORD_PATTERNS:
        hit = False
        for m in pat.finditer(stem):
            hit = True
            if m.end() == len(stem):  # keyword ends the stem -> names the contents
                return True
        if hit and len([w for w in _WORD_SPLIT.split(stem) if w]) <= 2:
            return True  # short name like token_config.yaml / secret_handler.txt
    return False

# Signals that a .md/.txt file is actually a converted academic paper
_PAPER_SIGNALS = [
    re.compile(r'\barxiv\b', re.IGNORECASE),
    re.compile(r'\bdoi\s*:', re.IGNORECASE),
    re.compile(r'\babstract\b', re.IGNORECASE),
    re.compile(r'\bproceedings\b', re.IGNORECASE),
    re.compile(r'\bjournal\b', re.IGNORECASE),
    re.compile(r'\bpreprint\b', re.IGNORECASE),
    re.compile(r'\\cite\{'),          # LaTeX citation
    re.compile(r'\[\d+\]'),           # Numbered citation [1], [23] (inline)
    re.compile(r'\[\n\d+\n\]'),       # Numbered citation spread across lines (markdown conversion)
    re.compile(r'eq\.\s*\d+|equation\s+\d+', re.IGNORECASE),
    re.compile(r'\d{4}\.\d{4,5}'),   # arXiv ID like 1706.03762
    re.compile(r'\bwe propose\b', re.IGNORECASE),   # common academic phrasing
    re.compile(r'\bliterature\b', re.IGNORECASE),   # "from the literature"
]
_PAPER_SIGNAL_THRESHOLD = 3  # need at least this many signals to call it a paper


def _is_sensitive(path: Path) -> bool:
    """Return True if this file likely contains secrets and should be skipped."""
    # Stage 1: any PARENT directory is a known secrets dir (parts[:-1] excludes
    # the filename itself so a root-level file named "credentials" is not falsely
    # skipped — the name patterns in Stage 2 handle the filename).
    if any(part in _SENSITIVE_DIRS for part in path.parts[:-1]):
        return True
    # Stage 2: filename pattern match
    name = path.name
    if any(p.search(name) for p in _SENSITIVE_PATTERNS):
        return True
    # Stage 3: generic keywords, only when load-bearing in the name
    return _generic_keyword_hit(name)


def _looks_like_paper(path: Path) -> bool:
    """Heuristic: does this text file read like an academic paper?"""
    try:
        # Only scan first 3000 chars for speed
        text = path.read_text(encoding="utf-8", errors="ignore")[:3000]
        hits = sum(1 for pattern in _PAPER_SIGNALS if pattern.search(text))
        return hits >= _PAPER_SIGNAL_THRESHOLD
    except Exception:
        return False


_ASSET_DIR_MARKERS = {".imageset", ".xcassets", ".appiconset", ".colorset", ".launchimage"}


_SHEBANG_CODE_INTERPRETERS = {
    "python", "python3", "python2",
    "ruby", "perl", "node", "nodejs",
    "bash", "sh", "dash", "zsh", "fish", "ksh", "tcsh",
    "lua", "php", "julia", "Rscript",
}


def _split_env_s(value: str, rest: list[str]) -> list[str]:
    """Re-tokenize an `env -S`/`--split-string` packed command, prepending the
    operand to any trailing args. Returns the unpacked argv."""
    packed = " ".join([value, *rest]).strip()
    return shlex.split(packed)


def _env_command_args(args: list[str], *, allow_split: bool = True) -> list[str]:
    """Strip leading env(1) options and var assignments, return the trailing
    command argv. Covers macOS/BSD and GNU coreutils env documented spellings.

    POSIX/macOS short forms:
        env [-0iv] [-C workdir] [-P utilpath] [-S string]
            [-u name] [name=value ...] [utility [argument ...]]

    GNU coreutils long/compact forms additionally supported:
        --argv0=ARG / -a ARG / -aARG
        --unset=NAME / --unset NAME / -u NAME / -uNAME
        --chdir=DIR / --chdir DIR / -C DIR / -CDIR
        --split-string=STRING / --split-string STRING
        -S STRING / -SSTRING / -vS STRING / -vSSTRING
        --ignore-environment / --null / --debug / --list-signal-handling
        --default-signal[=SIG] / --ignore-signal[=SIG] / --block-signal[=SIG]

    `-S` / `--split-string` payloads are themselves env-style argument lists
    per the GNU shebang synopsis:
        #!/usr/bin/env -[v]S[option]... [name=value]... command [args]...
    so after splitting the payload we recursively re-parse it with
    `allow_split=False` (a nested -S inside a split payload is rejected to
    bound recursion).

    Unknown hyphen-prefixed args yield [] (we refuse to guess whether
    their next token is an interpreter or an operand).
    """
    i = 0
    while i < len(args):
        arg = args[i]

        if arg == "--":
            return args[i + 1:]

        # Split-string forms: tokenize the packed payload, then re-parse it
        # as env args (so leading assignments/flags inside the payload are
        # skipped before the interpreter is identified).
        if allow_split:
            if arg == "-S":
                if i + 1 >= len(args):
                    return []
                return _env_command_args(
                    _split_env_s(" ".join(args[i + 1:]), []),
                    allow_split=False,
                )
            if arg.startswith("-S") and len(arg) > 2:
                return _env_command_args(
                    _split_env_s(arg[2:], args[i + 1:]),
                    allow_split=False,
                )
            if arg == "-vS":
                if i + 1 >= len(args):
                    return []
                return _env_command_args(
                    _split_env_s(" ".join(args[i + 1:]), []),
                    allow_split=False,
                )
            if arg.startswith("-vS") and len(arg) > 3:
                return _env_command_args(
                    _split_env_s(arg[3:], args[i + 1:]),
                    allow_split=False,
                )
            if arg.startswith("--split-string="):
                return _env_command_args(
                    _split_env_s(arg.split("=", 1)[1], args[i + 1:]),
                    allow_split=False,
                )
            if arg == "--split-string":
                if i + 1 >= len(args):
                    return []
                return _env_command_args(
                    _split_env_s(args[i + 1], args[i + 2:]),
                    allow_split=False,
                )

        # Options with separate required operand
        if arg in {"-u", "-C", "-P", "-a", "--unset", "--chdir", "--argv0"}:
            if i + 2 > len(args):
                return []
            i += 2
            continue

        # Clumped short option + operand
        if (
            arg.startswith(("-u", "-C", "-P", "-a"))
            and len(arg) > 2
            and not arg.startswith("--")
        ):
            i += 1
            continue

        # Long option with `=` operand
        if arg.startswith(("--unset=", "--chdir=", "--argv0=")):
            i += 1
            continue

        # No-operand flags
        if arg in {"-", "-i", "-0", "-v", "--ignore-environment", "--null",
                   "--debug", "--list-signal-handling"}:
            i += 1
            continue

        # Signal-handling long flags (with or without =SIG operand — we treat
        # them as no-effect for interpreter-resolution purposes)
        if arg.startswith(("--default-signal", "--ignore-signal", "--block-signal")):
            i += 1
            continue

        # Unknown hyphen-prefixed: refuse to guess
        if arg.startswith("-"):
            return []

        # Inline NAME=value assignment
        if "=" in arg:
            i += 1
            continue

        # First non-option, non-assignment token starts the command argv
        return args[i:]

    return []


def _shebang_interpreter(path: Path) -> str | None:
    """Return the interpreter name from a shebang line.

    Handles forms that a naive parser misses:
      - `#!/usr/bin/env -S python3 -u`     (env -S split-args form, anywhere)
      - `#!/usr/bin/env -i bash`           (no-operand env flags)
      - `#!/usr/bin/env -u VAR python3`    (env options with operands)
      - `#!/usr/bin/env -C /tmp python3`   (env -C workdir)
      - `#!/usr/bin/env -P /bin python3`   (env -P utilpath)
      - `#!/usr/bin/env DEBUG=1 python3`   (inline var assignment)
      - `#!"/usr/local/bin/python with spaces"`  (shlex handles quotes)

    Returns the basename of the resolved interpreter, or None if there is
    no shebang / the file is unreadable / parsing fails.
    """
    try:
        with path.open("rb") as f:
            first = f.read(256)
        if not first.startswith(b"#!"):
            return None
        line = first.split(b"\n")[0].decode(errors="replace")[2:].strip()
        parts = shlex.split(line)
        if not parts:
            return None
        interp = Path(parts[0]).name
        if interp == "env":
            env_args = _env_command_args(parts[1:])
            if not env_args:
                return None
            interp = Path(env_args[0]).name
        return interp
    except (OSError, ValueError):
        return None


def _shebang_file_type(path: Path) -> FileType | None:
    """Peek at the first line of an extensionless file for a shebang."""
    interp = _shebang_interpreter(path)
    if interp in _SHEBANG_CODE_INTERPRETERS:
        return FileType.CODE
    return None


def classify_file(path: Path) -> FileType | None:
    # Package manifests (apm.yml, pyproject.toml, go.mod, pom.xml) are parsed
    # deterministically, so route them to the AST path (CODE) rather than the LLM
    # document path — otherwise apm.yml (a .yml "document") would be LLM-extracted
    # and a package would split into duplicate file-anchored nodes (#1377).
    from graphify.manifest_ingest import is_package_manifest_path
    if is_package_manifest_path(path):
        return FileType.CODE
    # Compound extensions must be checked before simple suffix lookup
    if path.name.lower().endswith(".blade.php"):
        return FileType.CODE
    ext = path.suffix.lower()
    if not ext:
        return _shebang_file_type(path)
    if ext in CODE_EXTENSIONS:
        return FileType.CODE
    if ext in PAPER_EXTENSIONS:
        # PDFs inside Xcode asset catalogs are vector icons, not papers
        if any(part.endswith(tuple(_ASSET_DIR_MARKERS)) for part in path.parts):
            return None
        return FileType.PAPER
    if ext in IMAGE_EXTENSIONS:
        return FileType.IMAGE
    if ext in DOC_EXTENSIONS:
        # Check if it's a converted paper
        if _looks_like_paper(path):
            return FileType.PAPER
        return FileType.DOCUMENT
    if ext in OFFICE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in GOOGLE_WORKSPACE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in VIDEO_EXTENSIONS:
        return FileType.VIDEO
    return None


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF file using pypdf."""
    if not _file_within_size_cap(path):
        return ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx."""
    if not _zip_within_caps(path):
        return ""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    if not _zip_within_caps(path):
        return ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_extract_structure(path: Path) -> dict:
    """Extract structural nodes (sheets, named tables, column headers) from an .xlsx file.

    Returns a nodes/edges dict compatible with the graphify extract pipeline.
    Used in addition to xlsx_to_markdown so Claude sees both structure and content.
    """
    def _nid(*parts: str) -> str:
        return re.sub(r"[^a-z0-9_]", "_", "_".join(p.lower() for p in parts).strip("_"))

    try:
        import openpyxl
    except ImportError:
        return {"nodes": [], "edges": []}

    try:
        wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
    except Exception:
        return {"nodes": [], "edges": []}

    # F-035: typo fix — was `_re.sub` (NameError, but unreachable because the
    # whole xlsx codepath is currently behind a feature flag / not yet wired
    # into the dispatcher). Before re-enabling this path, re-audit it for
    # zip/XML bombs (openpyxl is built on top of zipfile and lxml-style XML
    # parsing — a malicious .xlsx can blow up memory at load_workbook time).
    stem = re.sub(r"[^a-z0-9]", "_", path.stem.lower())
    str_path = str(path)
    file_nid = _nid(str_path)
    nodes: list[dict] = [{"id": file_nid, "label": path.name, "file_type": "document",
                           "source_file": str_path, "source_location": None}]
    edges: list[dict] = []
    seen: set[str] = {file_nid}

    def _add(nid: str, label: str) -> None:
        if nid not in seen:
            seen.add(nid)
            nodes.append({"id": nid, "label": label, "file_type": "document",
                           "source_file": str_path, "source_location": None})

    def _edge(src: str, tgt: str, relation: str) -> None:
        edges.append({"source": src, "target": tgt, "relation": relation,
                       "confidence": "EXTRACTED", "source_file": str_path,
                       "source_location": None, "weight": 1.0})

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_nid = _nid(stem, sheet_name)
        _add(sheet_nid, f"{sheet_name} (sheet)")
        _edge(file_nid, sheet_nid, "contains")

        # Named Excel Tables (ListObjects)
        if hasattr(ws, "tables"):
            for tbl in ws.tables.values():
                tbl_nid = _nid(stem, sheet_name, tbl.name)
                _add(tbl_nid, tbl.name)
                _edge(sheet_nid, tbl_nid, "contains")
                # Column headers from table header row
                ref = tbl.ref  # e.g. "A1:D10"
                if ref:
                    try:
                        from openpyxl.utils import range_boundaries
                        min_col, min_row, max_col, _ = range_boundaries(ref)
                        header_row = list(ws.iter_rows(min_row=min_row, max_row=min_row,
                                                       min_col=min_col, max_col=max_col,
                                                       values_only=True))
                        if header_row:
                            for col_name in header_row[0]:
                                if col_name:
                                    col_nid = _nid(stem, tbl.name, str(col_name))
                                    _add(col_nid, str(col_name))
                                    _edge(tbl_nid, col_nid, "contains")
                    except Exception:
                        pass
        else:
            # Fallback: first non-empty row as column headers
            for row in ws.iter_rows(max_row=1, values_only=True):
                for cell in row:
                    if cell:
                        col_nid = _nid(stem, sheet_name, str(cell))
                        _add(col_nid, str(cell))
                        _edge(sheet_nid, col_nid, "contains")
                break

    try:
        wb.close()
    except Exception:
        pass

    return {"nodes": nodes, "edges": edges}


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # Use a stable name derived from the original path to avoid collisions.
    # Normalize the resolved path to NFC before hashing: on macOS (HFS+/APFS)
    # os.walk/rglob return filenames in NFD, while Python string literals and
    # directly-constructed Path objects are NFC, so the same source file would
    # otherwise hash to different sidecar names across runs — causing --update
    # to treat every Office file as new and re-extract it (#1226).
    import hashlib
    import unicodedata
    normalized_path = unicodedata.normalize("NFC", str(path.resolve()))
    name_hash = hashlib.sha256(normalized_path.encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    # Once the hash is stable the sidecar name is deterministic; skip re-writing
    # an existing sidecar so an unchanged source never churns its mtime (which
    # would still flag it as changed in detect_incremental).
    if out_path.exists():
        return out_path
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path


def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        return len(path.read_text(encoding="utf-8", errors="ignore").split())
    except Exception:
        return 0


# Directory names to always skip - venvs, caches, build artifacts, deps
_SKIP_DIRS = {
    "venv", ".venv", "env", ".env",
    "node_modules", "__pycache__", ".git",
    "dist", "build", "target", "out",
    "site-packages", "lib64",
    ".pytest_cache", ".mypy_cache", ".ruff_cache",
    ".tox", ".eggs", "*.egg-info",
    "graphify-out", GRAPHIFY_OUT_NAME,  # never treat own output as source input (#524); honour GRAPHIFY_OUT (#1423)
    # Coverage/test-artefact dirs — generated, never architecturally meaningful
    "coverage", "lcov-report",              # Vitest/Istanbul/nyc HTML reports (#870)
    "visual-tests", "visual-test",          # Playwright/visual-regression bundles (#869)
    "__snapshots__", "snapshots",           # Jest/Vitest snapshot dirs
    "storybook-static",                     # Storybook production build output
    "dist-protected",                       # Protected dist variants (same noise as dist)
    # Framework cache/build dirs — generated, never architecturally meaningful (#873)
    ".next", ".nuxt", ".turbo", ".angular",
    ".idea", ".cache", ".parcel-cache", ".svelte-kit", ".terraform", ".serverless",
    ".graphify",  # graphify's own extraction cache — never index self-generated data
    ".worktrees",  # git worktree convention (#947) — sibling checkouts, always redundant
}

# Large generated files that are never useful to extract
_SKIP_FILES = {
    "package-lock.json", "yarn.lock", "pnpm-lock.yaml",
    "Cargo.lock", "poetry.lock", "Gemfile.lock",
    "composer.lock", "go.sum", "go.work.sum",
}

def _is_noise_dir(part: str, parent: "Path | None" = None) -> bool:
    """Return True if this directory name looks like a venv, cache, or dep dir."""
    if part in _SKIP_DIRS:
        return True
    # Catch *_venv, *_repo/site-packages patterns
    if part.endswith("_venv") or part.endswith("_env"):
        return True
    if part.endswith(".egg-info"):
        return True
    # worktrees/ nested inside a dotted dir (e.g. .claude/worktrees/, .git/worktrees/)
    if part == "worktrees" and parent is not None and parent.name.startswith("."):
        return True
    return False


_VCS_MARKERS = (".git", ".hg", ".svn", "_darcs", ".fossil")


def _parse_gitignore_line(raw: str) -> str:
    """Parse one raw line from a .graphifyignore file per gitignore spec.

    - Strip newline chars
    - Strip inline comments (whitespace + # suffix), but only when # is
      preceded by whitespace — so path#with#hash.py is preserved
    - Unescape \\# to literal #
    - Remove trailing spaces unless escaped with backslash
    - Strip leading whitespace
    - Return empty string for blank lines and full-line comments
    """
    line = raw.rstrip("\n\r")
    line = line.lstrip()
    if not line or line.startswith("#"):
        return ""
    # Strip inline comments: require whitespace before # (gitignore extension)
    line = re.sub(r"\s+#+[^\\].*$", "", line)
    # Unescape \# → literal #
    line = line.replace("\\#", "#")
    # Remove unescaped trailing spaces (per gitignore spec)
    line = re.sub(r"(?<!\\) +$", "", line)
    return line


def _find_vcs_root(start: Path) -> Path | None:
    """Walk upward from start; return the first directory containing a VCS marker."""
    current = start.resolve()
    home = Path.home()
    while True:
        if any((current / m).exists() for m in _VCS_MARKERS):
            return current
        parent = current.parent
        if parent == current or current == home:
            return None
        current = parent


def _load_graphifyignore(root: Path) -> list[tuple[Path, str]]:
    """Read .graphifyignore files and return (anchor_dir, pattern) pairs.

    Patterns are returned outer-first so that inner (closer) rules are
    appended last and win via last-match-wins semantics — matching gitignore
    behavior exactly.

    Walk ceiling: the nearest VCS root if inside a repo, otherwise the scan
    root itself (hermetic — no leakage across unrelated sibling projects).
    """
    root = root.resolve()
    ceiling = _find_vcs_root(root) or root

    # Collect ancestor dirs from ceiling down to root (outer → inner)
    dirs: list[Path] = []
    current = root
    while True:
        dirs.append(current)
        if current == ceiling:
            break
        current = current.parent
    dirs.reverse()  # ceiling first, scan root last

    patterns: list[tuple[Path, str]] = []
    for d in dirs:
        # Merge .gitignore and .graphifyignore for this dir (#1363). Previously
        # the presence of a .graphifyignore made graphify skip that dir's
        # .gitignore entirely, so a file excluded only by .gitignore (e.g. a
        # neutrally-named secret like prod-dump.sql) silently got indexed into
        # the graph — whose artifacts embed file contents and are often
        # committed. .gitignore is read first and .graphifyignore last, so
        # .graphifyignore patterns (including `!` negations) win on conflict via
        # last-match-wins; adding a .graphifyignore can only ever exclude MORE,
        # never re-include a .gitignore-excluded file (#945 kept: a project with
        # only a .gitignore still gets sensible defaults).
        for fname in (".gitignore", ".graphifyignore"):
            ignore_file = d / fname
            if ignore_file.exists():
                for raw in ignore_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                    line = _parse_gitignore_line(raw)
                    if line:
                        patterns.append((d, line))
    return patterns


def _is_ignored(
    path: Path,
    root: Path,
    patterns: list[tuple[Path, str]],
    *,
    _cache: dict[Path, bool] | None = None,
) -> bool:
    """Return True if the path should be ignored per .graphifyignore patterns.

    Uses gitignore last-match-wins semantics: all patterns are evaluated in
    order; the final matching pattern determines the result. Negation patterns
    (starting with !) un-ignore a previously ignored path.

    Enforces gitignore's parent-exclusion rule: a ! pattern cannot re-include
    a file whose ancestor directory is already excluded.

    _cache: optional dict shared across calls within the same scan. Ancestor
    directory results are memoised so files under the same subtree don't
    re-evaluate the same patterns repeatedly.
    """
    if not patterns:
        return False

    def _eval(target: Path) -> bool:
        """Apply last-match-wins to a single target path."""
        if _cache is not None and target in _cache:
            return _cache[target]
        def _matches(rel: str, p: str, anchored: bool) -> bool:
            if anchored:
                return fnmatch.fnmatch(rel, p)
            parts = rel.split("/")
            if fnmatch.fnmatch(rel, p):
                return True
            if fnmatch.fnmatch(target.name, p):
                return True
            for i, part in enumerate(parts):
                if fnmatch.fnmatch(part, p):
                    return True
                if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                    return True
            return False

        result = False
        for anchor, pattern in patterns:
            negated = pattern.startswith("!")
            raw = pattern[1:] if negated else pattern
            anchored = raw.startswith("/")
            p = raw.strip("/")
            if not p:
                continue

            matched = False
            if anchored:
                try:
                    rel_anchor = str(target.relative_to(anchor)).replace(os.sep, "/")
                    matched = _matches(rel_anchor, p, anchored=True)
                except ValueError:
                    pass
            else:
                try:
                    rel = str(target.relative_to(root)).replace(os.sep, "/")
                    matched = _matches(rel, p, anchored=False)
                except ValueError:
                    pass
                if not matched and anchor != root:
                    try:
                        rel_anchor = str(target.relative_to(anchor)).replace(os.sep, "/")
                        matched = _matches(rel_anchor, p, anchored=False)
                    except ValueError:
                        pass

            if matched:
                result = not negated  # last match wins; ! flips to un-ignore
        if _cache is not None:
            _cache[target] = result
        return result

    # Gitignore parent-exclusion rule: a ! re-include cannot rescue a file
    # whose ancestor directory is already excluded. Walk ancestors top-down;
    # if any ancestor is excluded, the file is excluded regardless of later
    # ! patterns targeting the file or a sub-path.
    try:
        rel_parts = path.relative_to(root).parts
    except ValueError:
        return _eval(path)

    ancestor = root
    for part in rel_parts[:-1]:
        ancestor = ancestor / part
        if _eval(ancestor):
            return True
    return _eval(path)


def _load_graphifyinclude(root: Path) -> list[tuple[Path, str]]:
    """Read .graphifyinclude allowlist patterns from root and ancestors.

    Include patterns opt matching hidden files/dirs into traversal. Sensitive
    files and hard-skipped noise directories are still excluded later.
    Uses the same VCS-root ceiling logic as _load_graphifyignore.
    """
    root = root.resolve()
    ceiling = _find_vcs_root(root) or root

    dirs: list[Path] = []
    current = root
    while True:
        dirs.append(current)
        if current == ceiling:
            break
        current = current.parent
    dirs.reverse()

    patterns: list[tuple[Path, str]] = []
    for d in dirs:
        include_file = d / ".graphifyinclude"
        if include_file.exists():
            for raw in include_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                line = _parse_gitignore_line(raw)
                if line:
                    patterns.append((d, line))
    return patterns


def _is_included(path: Path, root: Path, patterns: list[tuple[Path, str]]) -> bool:
    """Return True if path matches any .graphifyinclude allowlist pattern."""
    if not patterns:
        return False

    def _matches(rel: str, p: str, anchored: bool) -> bool:
        if anchored:
            return fnmatch.fnmatch(rel, p)
        parts = rel.split("/")
        if fnmatch.fnmatch(rel, p):
            return True
        if fnmatch.fnmatch(path.name, p):
            return True
        for i, part in enumerate(parts):
            if fnmatch.fnmatch(part, p):
                return True
            if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                return True
        return False

    for anchor, pattern in patterns:
        anchored = pattern.startswith("/")
        p = pattern.strip("/")
        if not p:
            continue
        if anchored:
            try:
                rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                if _matches(rel_anchor, p, anchored=True):
                    return True
            except ValueError:
                pass
        else:
            try:
                rel = str(path.relative_to(root)).replace(os.sep, "/")
                if _matches(rel, p, anchored=False):
                    return True
            except ValueError:
                pass
            if anchor != root:
                try:
                    rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                    if _matches(rel_anchor, p, anchored=False):
                        return True
                except ValueError:
                    pass
    return False


def _could_contain_included_path(path: Path, root: Path, patterns: list[tuple[Path, str]]) -> bool:
    """Return True if a directory may contain files matched by .graphifyinclude."""
    if not patterns:
        return False

    rels: list[str] = []
    try:
        rels.append(str(path.relative_to(root)).replace(os.sep, "/"))
    except ValueError:
        pass
    for anchor, _ in patterns:
        if anchor != root:
            try:
                rels.append(str(path.relative_to(anchor)).replace(os.sep, "/"))
            except ValueError:
                pass

    for rel in rels:
        rel = rel.strip("/")
        if not rel:
            return True
        for _, pattern in patterns:
            p = pattern.strip("/")
            if not p:
                continue
            if p == rel or p.startswith(rel + "/"):
                return True
            if fnmatch.fnmatch(rel, p):
                return True
    return False


def _auto_follow_symlinks(root: Path) -> bool:
    """Auto-detect: ``True`` if ``root`` has any direct symlinked child.

    Allows "fake working dir" patterns (e.g. a folder full of symlinks pointing
    at scattered source dirs across the user's machine) to work transparently
    without the caller having to know to pass ``follow_symlinks=True``.

    Override is always possible by passing an explicit ``follow_symlinks=True``
    or ``follow_symlinks=False`` to :func:`detect` / :func:`detect_incremental`.
    """
    try:
        for p in root.iterdir():
            if p.is_symlink():
                return True
    except (OSError, PermissionError):
        pass
    return False


def detect(root: Path, *, follow_symlinks: bool | None = None, google_workspace: bool | None = None, extra_excludes: list[str] | None = None) -> dict:
    root = root.resolve()
    if follow_symlinks is None:
        follow_symlinks = _auto_follow_symlinks(root)
    google_workspace = google_workspace_enabled() if google_workspace is None else google_workspace
    files: dict[FileType, list[str]] = {
        FileType.CODE: [],
        FileType.DOCUMENT: [],
        FileType.PAPER: [],
        FileType.IMAGE: [],
        FileType.VIDEO: [],
    }
    total_words = 0

    skipped_sensitive: list[str] = []
    ignore_patterns = _load_graphifyignore(root)
    ignore_cache: dict[Path, bool] = {}  # shared across all _is_ignored calls in this scan
    # CLI --exclude patterns are anchored at the scan root and appended last
    # so they win over any .graphifyignore/.gitignore rules (#947).
    if extra_excludes:
        for pat in extra_excludes:
            line = _parse_gitignore_line(pat)
            if line:
                ignore_patterns.append((root, line))
    include_patterns = _load_graphifyinclude(root)

    # Always include graphify-out/memory/ - query results filed back into the graph
    memory_dir = root / GRAPHIFY_OUT / "memory"
    scan_paths = [root]
    if memory_dir.exists():
        scan_paths.append(memory_dir)

    seen: set[Path] = set()
    all_files: list[Path] = []

    for scan_root in scan_paths:
        in_memory_tree = memory_dir.exists() and str(scan_root).startswith(str(memory_dir))
        for dirpath, dirnames, filenames in os.walk(scan_root, followlinks=follow_symlinks):
            dp = Path(dirpath)
            if follow_symlinks and os.path.islink(dirpath):
                real = os.path.realpath(dirpath)
                parent_real = os.path.realpath(os.path.dirname(dirpath))
                if parent_real == real or parent_real.startswith(real + os.sep):
                    dirnames.clear()
                    continue
            if not in_memory_tree:
                # Prune noise dirs in-place so os.walk never descends into them.
                # Dot dirs are allowed — users often want .github/, .claude/, etc.
                # Framework caches (.next, .nuxt, …) are caught by _is_noise_dir.
                # Negations need no special-casing here: _is_ignored already applies
                # last-match-wins (so `!dir/` un-ignores a directory and it won't be
                # pruned) and the gitignore parent-exclusion rule (a `!` cannot rescue
                # a file beneath an excluded dir), so descending an ignored directory to
                # look for a re-included file is never necessary. The previous blanket
                # `has_negation` disabled directory pruning for EVERY ignored dir whenever
                # any `!` rule existed — e.g. a single `!docs/**` made the walk descend
                # bin/, obj/, wwwroot/, generated/, … : a pathological slowdown on large
                # repos for no correctness gain.
                dirnames[:] = [
                    d for d in dirnames
                    if not _is_noise_dir(d, dp)
                    and not _is_ignored(dp / d, root, ignore_patterns, _cache=ignore_cache)
                ]
            for fname in filenames:
                if fname in _SKIP_FILES:
                    continue
                p = dp / fname
                if p not in seen:
                    seen.add(p)
                    all_files.append(p)

    all_files.sort(key=lambda p: str(p))

    converted_dir = root / GRAPHIFY_OUT / "converted"

    for p in all_files:
        # For memory dir files, skip hidden/noise filtering
        in_memory = memory_dir.exists() and str(p).startswith(str(memory_dir))
        if not in_memory:
            # Skip files inside our own converted/ dir (avoid re-processing sidecars)
            if str(p).startswith(str(converted_dir)):
                continue
        if not in_memory and _is_ignored(p, root, ignore_patterns, _cache=ignore_cache):
            continue
        if _is_sensitive(p):
            skipped_sensitive.append(str(p))
            continue
        ftype = classify_file(p)
        if ftype:
            if p.suffix.lower() in GOOGLE_WORKSPACE_EXTENSIONS:
                if not google_workspace:
                    skipped_sensitive.append(
                        str(p)
                        + " [Google Workspace shortcut skipped - pass --google-workspace "
                        "or set GRAPHIFY_GOOGLE_WORKSPACE=1]"
                    )
                    continue
                try:
                    md_path = convert_google_workspace_file(p, converted_dir, xlsx_to_markdown=xlsx_to_markdown)
                except Exception as exc:
                    skipped_sensitive.append(str(p) + f" [Google Workspace export failed: {exc}]")
                    continue
                if md_path:
                    if _is_ignored(md_path, root, ignore_patterns, _cache=ignore_cache):
                        continue
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    skipped_sensitive.append(str(p) + " [Google Workspace export produced no readable text]")
                continue
            # Office files: convert to markdown sidecar so subagents can read them
            if p.suffix.lower() in OFFICE_EXTENSIONS:
                md_path = convert_office_file(p, converted_dir)
                if md_path:
                    if _is_ignored(md_path, root, ignore_patterns, _cache=ignore_cache):
                        continue
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    # Conversion failed (library not installed) - skip with note
                    skipped_sensitive.append(str(p) + " [office conversion failed - pip install graphifyy[office]]")
                continue
            files[ftype].append(str(p))
            if ftype != FileType.VIDEO:
                total_words += count_words(p)

    for ftype in files:
        files[ftype].sort()

    total_files = sum(len(v) for v in files.values())
    needs_graph = total_words >= CORPUS_WARN_THRESHOLD

    # Determine warning - lower bound, upper bound, or sensitive files skipped
    warning: str | None = None
    if not needs_graph:
        warning = (
            f"Corpus is ~{total_words:,} words - fits in a single context window. "
            f"You may not need a graph."
        )
    elif total_words >= CORPUS_UPPER_THRESHOLD or total_files >= FILE_COUNT_UPPER:
        warning = (
            f"Large corpus: {total_files} files · ~{total_words:,} words. "
            f"Semantic extraction will be expensive (many Claude tokens). "
            f"Consider running on a subfolder."
        )

    return {
        "files": {k.value: v for k, v in files.items()},
        "total_files": total_files,
        "total_words": total_words,
        "needs_graph": needs_graph,
        "warning": warning,
        "skipped_sensitive": skipped_sensitive,
        "graphifyignore_patterns": len(ignore_patterns),
        "scan_root": str(root.resolve()),
    }


def _md5_file(path: Path) -> str:
    """MD5 of file contents streamed in 64KB chunks — for change detection only."""
    import hashlib as _hl
    h = _hl.md5(usedforsecurity=False)
    try:
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
    except OSError:
        return ""
    return h.hexdigest()


def _stat_and_hash(path_str: str) -> tuple[str, float, str] | None:
    """Stat + MD5 a single file; returns None on OSError (e.g. deleted mid-run)."""
    try:
        p = Path(path_str)
        return path_str, p.stat().st_mtime, _md5_file(p)
    except OSError:
        return None


def _to_relative_for_storage(key: str, root: Path) -> str:
    """Return ``key`` as a forward-slash relative path from ``root``.

    Keys outside ``root`` (out-of-tree symlinked sources, external --include
    paths) and already-relative keys pass through unchanged — mirrors the
    fallback in :func:`graphify.watch._relativize_source_files` so the
    on-disk artifact survives the round-trip even when some paths cannot be
    portably encoded.

    Only ``root`` is resolved — the key itself is relativized symbolically
    so an in-root symlink (e.g. ``alias.py -> sub/target.py``) is stored
    under its own name. Resolving the key would point the stored entry at
    the symlink target, and the original key would then miss on reload and
    re-extract on every incremental run.
    """
    p = Path(key)
    if not p.is_absolute():
        return key
    try:
        rel = os.path.relpath(p, Path(root).resolve())
    except (ValueError, OSError):
        return key  # outside root (e.g. Windows cross-drive)
    # ``os.path.relpath`` happily produces ``../foo`` for paths outside
    # root; mirror the prior ``relative_to``-raises-ValueError semantics by
    # keeping out-of-root entries in their absolute form.
    if rel == ".." or rel.startswith(".." + os.sep) or rel.startswith("../"):
        return key
    return rel.replace(os.sep, "/")


def _to_absolute_from_storage(key: str, root: Path) -> str:
    """Inverse of :func:`_to_relative_for_storage`.

    Re-anchor a stored key against ``root``. Already-absolute keys
    (legacy manifests, out-of-root entries) pass through unchanged so
    that newly-loaded manifests from before this change remain readable.
    Uses ``Path(root).resolve()`` so the produced absolute path matches
    what :func:`detect` returns (which also resolves the scan root).
    """
    p = Path(key)
    if p.is_absolute():
        return str(p)
    return str(Path(root).resolve() / p)


def load_manifest(
    manifest_path: str = _MANIFEST_PATH,
    *,
    root: Path | None = None,
) -> dict:
    """Load the manifest from a previous run. Returns {} on any error.

    When ``root`` is provided, stored relative keys are re-anchored against
    it so callers see absolute paths regardless of on-disk format. Legacy
    manifests with absolute keys pass through unchanged, so a graphify-out/
    written by an older version (or by a caller that didn't supply ``root``
    to :func:`save_manifest`) remains readable.
    """
    try:
        raw = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}
    if root is None or not isinstance(raw, dict):
        return raw
    return {_to_absolute_from_storage(k, root): v for k, v in raw.items()}


def save_manifest(
    files: dict[str, list[str]],
    manifest_path: str = _MANIFEST_PATH,
    *,
    kind: str = "both",
    root: Path | None = None,
) -> None:
    """Save current file mtimes + content hashes for change detection.

    kind="ast"      — written by `graphify update` (AST-only rebuild). Stamps
                      ast_hash; preserves an existing semantic_hash only when
                      the file content is unchanged (mtime + hash match).
    kind="semantic" — written by `graphify extract` after semantic extraction.
                      Stamps semantic_hash; preserves existing ast_hash.
    kind="both"     — full pipeline: stamps both hashes (default).

    When ``root`` is provided, keys are relativized against it before write
    (forward-slash, posix-style) so the on-disk manifest is portable across
    machines and checkout locations (#777). Out-of-root entries are written
    as absolute so they continue to round-trip on the saving machine.
    When ``root`` is None the legacy absolute-keyed format is preserved.
    """
    existing = load_manifest(manifest_path, root=root)

    def _normalise_entry(entry):
        if isinstance(entry, (int, float)):
            return {"mtime": entry, "ast_hash": "", "semantic_hash": ""}
        if isinstance(entry, dict) and "hash" in entry and "ast_hash" not in entry:
            return {"mtime": entry.get("mtime", 0), "ast_hash": entry["hash"], "semantic_hash": ""}
        if isinstance(entry, dict):
            return entry
        return None

    # Seed from the existing manifest so incremental callers passing a subset
    # of files don't silently erase entries for untouched files (#917).
    # Prune entries whose file no longer exists on disk — those are genuine
    # deletions that detect_incremental() should treat as gone.
    manifest: dict[str, dict] = {}
    for f, entry in existing.items():
        normalised = _normalise_entry(entry)
        if normalised is None:
            continue
        try:
            if Path(f).exists():
                manifest[f] = normalised
        except OSError:
            continue

    all_files = [f for file_list in files.values() for f in file_list]
    with ThreadPoolExecutor() as pool:
        raw = pool.map(_stat_and_hash, all_files)
    hashed: dict[str, tuple[float, str]] = {
        r[0]: (r[1], r[2]) for r in raw if r is not None
    }

    for f in all_files:
        if f not in hashed:
            continue  # file deleted between detect() and manifest write
        mtime, h = hashed[f]
        prev = _normalise_entry(existing.get(f, {})) or {}
        entry: dict = {"mtime": mtime}
        if kind in ("ast", "both"):
            entry["ast_hash"] = h
        else:
            entry["ast_hash"] = prev.get("ast_hash", "")
        if kind in ("semantic", "both"):
            entry["semantic_hash"] = h
        else:
            # Preserve semantic_hash only when content is unchanged
            entry["semantic_hash"] = prev.get("semantic_hash", "") if h == prev.get("ast_hash", "") else ""
        manifest[f] = entry
    if root is not None:
        # Persist in portable form: forward-slash relative paths. Keys outside
        # ``root`` (out-of-tree symlinked corpora, --include sources) keep
        # their absolute form so the manifest round-trips on the saving
        # machine even when not every entry can be portably encoded.
        manifest = {_to_relative_for_storage(k, root): v for k, v in manifest.items()}
    Path(manifest_path).parent.mkdir(parents=True, exist_ok=True)
    Path(manifest_path).write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def detect_incremental(
    root: Path,
    manifest_path: str = _MANIFEST_PATH,
    *,
    follow_symlinks: bool | None = None,
    google_workspace: bool | None = None,
    kind: str = "semantic",
    extra_excludes: list[str] | None = None,
) -> dict:
    """Like detect(), but returns only new or modified files since the last run.

    kind="semantic" (default for extract): a file is "changed" when its
        semantic_hash is missing or its content has changed since the last
        semantic extraction pass. Use this for `graphify extract` so that
        files touched by `graphify update` (AST-only) are re-extracted
        semantically.
    kind="ast": a file is "changed" when its ast_hash is missing or its
        content has changed. Use this for `graphify update`.

    Fast path: mtime unchanged + hash matches → unchanged (free, no disk IO
    beyond stat). Slow path: mtime bumped → compare MD5 against the relevant
    hash field before re-extracting.

    Backwards compatible with legacy manifests storing plain float mtime values
    or {mtime, hash} dicts (treated as ast_hash only; semantic_hash = miss).

    The ``follow_symlinks`` flag is forwarded to :func:`detect` so corpora that
    rely on symlinked sub-trees (e.g. a ``state_of_truth/`` symlink pointing to a
    directory outside the scan root) are scanned consistently between full and
    incremental runs. ``None`` (default) means auto-detect: ``True`` when ``root``
    contains at least one direct symlinked child, ``False`` otherwise.
    """
    full = detect(root, follow_symlinks=follow_symlinks, google_workspace=google_workspace, extra_excludes=extra_excludes)
    # Pass ``root`` so a manifest written with relative keys (post-#777) is
    # re-anchored to the absolute form the rest of this function compares
    # against. Legacy absolute-keyed manifests pass through unchanged.
    manifest = load_manifest(manifest_path, root=root)

    if not manifest:
        # No previous run - treat everything as new
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}

    for ftype, file_list in full["files"].items():
        for f in file_list:
            stored = manifest.get(f)
            try:
                current_mtime = Path(f).stat().st_mtime
            except Exception:
                current_mtime = 0

            # Legacy manifest: plain float value — treat as ast_hash only
            if isinstance(stored, (int, float)):
                changed = stored is None or current_mtime > stored
            elif isinstance(stored, dict):
                # Normalise legacy {mtime, hash} to new schema
                if "hash" in stored and "ast_hash" not in stored:
                    stored = {"mtime": stored.get("mtime", 0), "ast_hash": stored["hash"], "semantic_hash": ""}
                hash_key = "semantic_hash" if kind == "semantic" else "ast_hash"
                stored_hash = stored.get(hash_key, "")
                # Missing semantic_hash means update ran but extract hasn't — always re-extract
                if not stored_hash:
                    changed = True
                else:
                    stored_mtime = stored.get("mtime")
                    # Schema-drift guard (#1163): tolerate a nested {mtime: ...}
                    # dict or any non-numeric value without crashing.
                    if isinstance(stored_mtime, dict):
                        stored_mtime = stored_mtime.get("mtime")
                    if not isinstance(stored_mtime, (int, float)):
                        stored_mtime = None
                    if stored_mtime is None or current_mtime != stored_mtime:
                        # mtime bumped — verify with content hash before re-extracting
                        changed = _md5_file(Path(f)) != stored_hash
                    else:
                        changed = False
            else:
                changed = True  # unknown format, re-extract to be safe

            if changed:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    # Files in manifest that no longer exist - their cached nodes are now ghost nodes
    current_files = {f for flist in full["files"].values() for f in flist}
    deleted_files = [f for f in manifest if f not in current_files]

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    return full
